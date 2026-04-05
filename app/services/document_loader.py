import io
from pathlib import Path
from typing import List, Tuple
from PyPDF2 import PdfReader
from docx import Document
from app.core.logging import get_logger

logger = get_logger(__name__)


class DocumentLoader:
    """
    Loads PDF and TXT files.
    Returns: List of (filename, raw_text) tuples.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}

    def load_directory(self, folder_path: str) -> List[Tuple[str, str]]:
        """
        Load all supported files inside a directory.
        Returns a list of (filename, text).
        """
        folder = Path(folder_path)
        if not folder.exists():
            logger.warning(f"Data folder not found: {folder_path}")
            return []

        results = []
        for file_path in sorted(folder.iterdir()):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    text = self._load_file(file_path)
                    if text.strip():
                        results.append((file_path.name, text))
                        logger.info(f"Loaded: {file_path.name} ({len(text)} chars)")
                    else:
                        logger.warning(f"Empty file skipped: {file_path.name}")
                except Exception as e:
                    logger.error(f"Failed to load {file_path.name}: {e}")

        logger.info(f"Total documents loaded: {len(results)}")
        return results

    def load_single_file(self, file_path: str) -> Tuple[str, str]:
        """Load a single file. Returns (filename, text)."""
        path = Path(file_path)
        text = self._load_file(path)
        return path.name, text

    def load_bytes(self, filename: str, content: bytes) -> Tuple[str, str]:
        """
        Load a document from in-memory bytes (used for uploads).
        Returns (filename, text).
        """
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            text = self._extract_pdf_bytes(content)
        elif suffix == ".txt":
            text = content.decode("utf-8", errors="ignore")
        elif suffix == ".docx":
            text = self._extract_docx_bytes(content)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
        return filename, text

    def _load_file(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._extract_pdf(path)
        elif suffix == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore")
        elif suffix == ".docx":
            return self._extract_docx(path)
        else:
            raise ValueError(f"Unsupported: {suffix}")

    def _extract_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    def _extract_pdf_bytes(self, content: bytes) -> str:
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    def _extract_docx(self, path: Path) -> str:
        doc = Document(str(path))
        return self._parse_docx_paragraphs(doc)

    def _extract_docx_bytes(self, content: bytes) -> str:
        doc = Document(io.BytesIO(content))
        return self._parse_docx_paragraphs(doc)

    def _parse_docx_paragraphs(self, doc: Document) -> str:
        """
        Parse DOCX paragraphs and wrap bold text with `** **`.

        The following 3 cases are handled:
        1. The full paragraph is bold -> treat it as a heading -> `**text**`
        2. The paragraph has mixed bold/normal runs -> wrap inline bold text
        3. A normal paragraph -> keep as-is

        This allows the preprocessor's `HEADING_PATTERN` (`** heading **`) to work
        regardless of which DOCX file the user uploads.
        """
        paragraphs = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            # Case 1: The entire paragraph is bold (heading)
            if self._is_fully_bold(para):
                paragraphs.append(f"**{para.text.strip()}**")
                continue

            # Case 2: Mixed runs - wrap bold parts
            if self._has_any_bold(para):
                line = self._extract_mixed_runs(para)
                paragraphs.append(line)
                continue

            # Case 3: Normal paragraph
            paragraphs.append(para.text.strip())

        return "\n\n".join(paragraphs)

    def _is_fully_bold(self, para) -> bool:
        """
        Check whether all non-empty runs in a paragraph are bold.
        """
        runs = [r for r in para.runs if r.text.strip()]
        if not runs:
            return False
        return all(r.bold for r in runs)

    def _has_any_bold(self, para) -> bool:
        """Check whether the paragraph contains any bold run."""
        return any(r.bold and r.text.strip() for r in para.runs)

    def _extract_mixed_runs(self, para) -> str:
        """
        Handle a mixed paragraph by wrapping bold runs with `** **`.
        Example: "This is an **important** word."
        """
        parts = []
        for run in para.runs:
            if not run.text:
                continue
            if run.bold:
                parts.append(f"**{run.text}**")
            else:
                parts.append(run.text)
        return "".join(parts)
